import streamlit as st
import json
import io
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pypdf
from google import genai
from google.genai import types

# ---------------------------------------------------------
# 1. TRÍCH XUẤT TÀI LIỆU (PDF/DOCX/TXT)
# ---------------------------------------------------------
class DocumentParser:
    @staticmethod
    def parse_uploaded_files(uploaded_files):
        combined_text = ""
        for file in uploaded_files:
            file_bytes = file.read()
            filename = file.name.lower()
            combined_text += f"\n\n--- TÀI LIỆU: {file.name} ---\n"
            
            if filename.endswith('.pdf'):
                pdf_reader = pypdf.PdfReader(io.BytesIO(file_bytes))
                for page_num, page in enumerate(pdf_reader.pages):
                    extracted = page.extract_text()
                    if extracted:
                        combined_text += f"\n[Trang {page_num + 1}]\n" + extracted
            elif filename.endswith('.docx'):
                doc = docx.Document(io.BytesIO(file_bytes))
                for para in doc.paragraphs:
                    if para.text.strip():
                        combined_text += para.text + "\n"
            elif filename.endswith('.txt'):
                combined_text += file_bytes.decode("utf-8", errors="ignore")
        return combined_text

# ---------------------------------------------------------
# 2. XUẤT FILE WORD TỰ ĐỘNG CHUẨN 5512
# ---------------------------------------------------------
class WordExporter:
    @staticmethod
    def generate_docx(data):
        doc = docx.Document()

        # Định dạng lề A4 chuẩn Công văn 5512
        section = doc.sections[0]
        section.top_margin = Inches(0.79)
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(0.79)

        # Bảng Thông tin Trường & Giáo viên
        s_info = data.get("school_info", {})
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        
        p_left = table.cell(0, 0).paragraphs[0]
        p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_left.add_run(f"Trường: {s_info.get('school_name', '')}\nTổ: {s_info.get('department', '')}").bold = True

        p_right = table.cell(0, 1).paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_right.add_run(f"Họ và tên giáo viên:\n{s_info.get('teacher_name', '')}").bold = True

        # Tên bài dạy
        g_info = data.get("general_info", {})
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.space_before = Pt(12)
        p_title.paragraph_format.space_after = Pt(6)
        
        r_title = p_title.add_run(f"TÊN BÀI DẠY: {g_info.get('lesson_title', '').upper()}\n")
        r_title.bold = True
        r_title.font.size = Pt(14)
        
        r_sub = p_title.add_run(f"Môn học: {g_info.get('subject', '')}; Lớp: {g_info.get('grade', '')}\nThời gian thực hiện: ({g_info.get('duration', '')})")
        r_sub.italic = True

        # Hàm thêm đoạn văn bản chuẩn Times New Roman 13pt
        def add_p(text, bold=False, italic=False, space_after=3, space_before=0, indent=0):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(space_before)
            p.paragraph_format.space_after = Pt(space_after)
            p.paragraph_format.line_spacing = 1.15
            if indent > 0:
                p.paragraph_format.left_indent = Inches(indent)
            run = p.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(13)
            run.bold = bold
            run.italic = italic
            return p

        # I. MỤC TIÊU
        add_p("I. Mục tiêu", bold=True, space_before=6)
        objs = data.get("objectives", {})
        
        add_p("1. Kiến thức", bold=True, indent=0.2)
        for k in objs.get("knowledge", []):
            add_p(f"- {k}", indent=0.4)
            
        add_p("2. Năng lực", bold=True, indent=0.2)
        comps = objs.get("competencies", {})
        add_p("2.1. Năng lực đặc thù:", bold=True, indent=0.3)
        for c in comps.get("subject_specific", []):
            add_p(f"- {c}", indent=0.5)
            
        add_p("2.2. Năng lực chung:", bold=True, indent=0.3)
        for c in comps.get("general", []):
            add_p(f"- {c}", indent=0.5)
            
        add_p("2.3. Năng lực Số & AI tích hợp:", bold=True, indent=0.3)
        for c in comps.get("digital_ai", []):
            add_p(f"- {c}", indent=0.5)

        add_p("3. Phẩm chất", bold=True, indent=0.2)
        for q in objs.get("qualities", []):
            add_p(f"- {q}", indent=0.4)

        # II. THIẾT BỊ DẠY HỌC VÀ HỌC LIỆU
        add_p("II. THIẾT BỊ DẠY HỌC VÀ HỌC LIỆU", bold=True, space_before=6)
        eq = data.get("equipment", {})
        add_p("1. Giáo viên (GV)", bold=True, indent=0.2)
        for item in eq.get("teacher", []):
            add_p(f"- {item}", indent=0.4)
        add_p("2. Học sinh (HS)", bold=True, indent=0.2)
        for item in eq.get("student", []):
            add_p(f"- {item}", indent=0.4)

        # III. TIẾN TRÌNH DẠY HỌC
        add_p("III. TIẾN TRÌNH DẠY HỌC", bold=True, space_before=6)
        for act in data.get("activities", []):
            add_p(act.get("activity_name", "").upper(), bold=True, space_before=4)
            add_p("a) Mục tiêu:", bold=True, italic=True, indent=0.2)
            add_p(act.get("objective", ""), indent=0.4)
            add_p("b) Nội dung:", bold=True, italic=True, indent=0.2)
            add_p(act.get("content", ""), indent=0.4)
            add_p("c) Sản phẩm:", bold=True, italic=True, indent=0.2)
            add_p(act.get("product", ""), indent=0.4)
            add_p("d) Tổ chức thực hiện:", bold=True, italic=True, indent=0.2)
            
            steps = act.get("implementation", {})
            add_p("Bước 1: Chuyển giao nhiệm vụ", bold=True, indent=0.3)
            add_p(steps.get("step_1", ""), indent=0.5)
            add_p("Bước 2: Thực hiện nhiệm vụ", bold=True, indent=0.3)
            add_p(steps.get("step_2", ""), indent=0.5)
            add_p("Bước 3: Báo cáo, thảo luận", bold=True, indent=0.3)
            add_p(steps.get("step_3", ""), indent=0.5)
            add_p("Bước 4: Kết luận, nhận định", bold=True, indent=0.3)
            add_p(steps.get("step_4", ""), indent=0.5)

        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)
        return stream

# ---------------------------------------------------------
# 3. GIAO DIỆN WEB STREAMLIT
# ---------------------------------------------------------
st.set_page_config(page_title="Trợ lý Soạn KHBD 5512", page_icon="📘", layout="wide")

st.title("📘 TRỢ LÝ AI SOẠN KẾ HOẠCH BÀI DẠY (CÔNG VĂN 5512)")
st.caption("Ứng dụng dành cho Giáo viên THPT - GDPT 2018")

with st.sidebar:
    st.header("🔑 Cấu hình")
    api_key = st.text_input("Nhập Gemini API Key:", type="password")
    st.info("Lấy API Key miễn phí tại: aistudio.google.com")

# Khởi tạo dữ liệu
if "lesson_data" not in st.session_state:
    st.session_state.lesson_data = None

# TAB GIAO DIỆN
tab1, tab2, tab3 = st.tabs(["1. Thông tin & Tài liệu", "2. Chỉnh sửa KHBD", "3. Xuất file Word"])

with tab1:
    col1, col2 = st.columns(2)
    with col1:
        school = st.text_input("Trường THPT:", "THPT Nguyễn Văn Trỗi")
        dept = st.text_input("Tổ chuyên môn:", "Tổ Toán")
        teacher = st.text_input("Họ tên giáo viên:", "Dương Tấn Tiến")
    with col2:
        subject = st.text_input("Môn học:", "Toán học")
        grade = st.selectbox("Lớp:", ["Lớp 10", "Lớp 11", "Lớp 12"])
        title = st.text_input("Tên bài dạy:", "BÀI 5: GIÁ TRỊ LƯỢNG GIÁC CỦA MỘT GÓC TỪ 0° ĐẾN 180°")
        duration = st.text_input("Thời lượng:", "2 tiết")

    files = st.file_uploader("Tải lên SGK / Sách giáo viên (PDF/DOCX/TXT):", accept_multiple_files=True)

    if st.button("🤖 Bắt đầu AI Soạn Bài", type="primary"):
        if not api_key:
            st.error("Vui lòng nhập Gemini API Key ở góc trái màn hình!")
        elif not files:
            st.error("Vui lòng tải lên ít nhất 1 file SGK hoặc tài liệu tham khảo!")
        else:
            with st.spinner("AI đang phân tích SGK và xây dựng KHBD chuẩn 5512..."):
                text_content = DocumentParser.parse_uploaded_files(files)
                
                # Gọi Gemini API
                client = genai.Client(api_key=api_key)
                prompt = f"""
                Trường: {school}, Tổ: {dept}, GV: {teacher}
                Môn: {subject}, Lớp: {grade}, Bài: {title}, Thời lượng: {duration}
                
                NỘI DUNG SGK THAM CHIẾU:
                {text_content}
                
                Hãy soạn Kế hoạch bài dạy theo Công văn 5512 dựa trên SGK. Nghiêm cấm bịa đặt kiến thức không có trong SGK.
                Trả về JSON tuân thủ đúng cấu trúc.
                """
                
                schema = {
                    "type": "OBJECT",
                    "properties": {
                        "school_info": {"type": "OBJECT", "properties": {"school_name": {"type": "STRING"}, "department": {"type": "STRING"}, "teacher_name": {"type": "STRING"}}},
                        "general_info": {"type": "OBJECT", "properties": {"lesson_title": {"type": "STRING"}, "subject": {"type": "STRING"}, "grade": {"type": "STRING"}, "duration": {"type": "STRING"}}},
                        "objectives": {
                            "type": "OBJECT",
                            "properties": {
                                "knowledge": {"type": "ARRAY", "items": {"type": "STRING"}},
                                "competencies": {"type": "OBJECT", "properties": {"subject_specific": {"type": "ARRAY", "items": {"type": "STRING"}}, "general": {"type": "ARRAY", "items": {"type": "STRING"}}, "digital_ai": {"type": "ARRAY", "items": {"type": "STRING"}}}},
                                "qualities": {"type": "ARRAY", "items": {"type": "STRING"}}
                            }
                        },
                        "equipment": {"type": "OBJECT", "properties": {"teacher": {"type": "ARRAY", "items": {"type": "STRING"}}, "student": {"type": "ARRAY", "items": {"type": "STRING"}}}},
                        "activities": {
                            "type": "ARRAY",
                            "items": {
                                "type": "OBJECT",
                                "properties": {
                                    "activity_name": {"type": "STRING"},
                                    "objective": {"type": "STRING"},
                                    "content": {"type": "STRING"},
                                    "product": {"type": "STRING"},
                                    "implementation": {"type": "OBJECT", "properties": {"step_1": {"type": "STRING"}, "step_2": {"type": "STRING"}, "step_3": {"type": "STRING"}, "step_4": {"type": "STRING"}}}
                                }
                            }
                        }
                    }
                }
                
                response = client.models.generate_content(
                    model='gemini-2.5-pro',
                    contents=prompt,
                    config=types.GenerateContentConfig(
                        response_mime_type="application/json",
                        response_schema=schema,
                        temperature=0.2
                    )
                )
                
                st.session_state.lesson_data = json.loads(response.text)
                st.success("Tạo KHBD thành công! Hãy chuyển sang Tab 2 để chỉnh sửa.")

with tab2:
    if st.session_state.lesson_data:
        st.subheader("✏️ Bạn có thể sửa trực tiếp nội dung bài dạy ở đây:")
        data = st.session_state.lesson_data
        
        data["general_info"]["lesson_title"] = st.text_input("Tên bài dạy:", data["general_info"]["lesson_title"])
        
        st.markdown("#### I. Mục tiêu")
        k_text = "\n".join(data["objectives"]["knowledge"])
        new_k = st.text_area("Kiến thức (Mỗi dòng 1 ý):", k_text, height=100)
        data["objectives"]["knowledge"] = [x.strip() for x in new_k.split("\n") if x.strip()]

        st.session_state.lesson_data = data
    else:
        st.info("Chưa có dữ liệu. Vui lòng thực hiện ở Tab 1 trước.")

with tab3:
    if st.session_state.lesson_data:
        st.subheader("📄 Xuất file Word chuẩn Công văn 5512")
        docx_bytes = WordExporter.generate_docx(st.session_state.lesson_data)
        
        st.download_button(
            label="📥 TẢI FILE WORD (.DOCX) VỀ MÁY",
            data=docx_bytes,
            file_name=f"KHBD_{st.session_state.lesson_data['general_info']['lesson_title']}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary"
        )
    else:
        st.info("Chưa có dữ liệu để xuất file.")
